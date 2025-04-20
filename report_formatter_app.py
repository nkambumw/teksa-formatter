import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
import os
import win32com.client as win32

st.set_page_config(page_title="TekSA Report Management System", layout="wide")

st.markdown("""
    <style>
        html, body, [class*="css"] {
            font-size: 18px !important;
        }
        .big-title {
            font-size: 48px !important;
            font-weight: 900;
            color: #002f6c;
            margin-bottom: 0.2em;
        }
        .sub-title {
            font-size: 22px !important;
            color: #444;
            margin-bottom: 1.5em;
        }
        .info-box {
            font-size: 18px;
            color: #0b3d91;
            background-color: #e0f2ff;
            border-left: 6px solid #0077b6;
            padding: 16px 20px;
            border-radius: 8px;
            margin-top: 25px;
        }
        .stDownloadButton > button, .stButton > button {
            font-size: 18px !important;
            padding: 10px 22px;
        }
        .stFileUploader {
            font-size: 18px;
        }
    </style>
""", unsafe_allow_html=True)

st.sidebar.title("🛠 Formatting Options")

format_mode = st.sidebar.radio(
    "Choose formatting method:",
    ("📄 Upload a template", "⚙️ Choose on-the-fly options")
)

template_path = None
formatting_choices = {}

if format_mode == "📄 Upload a template":
    template_file = st.sidebar.file_uploader("📂 Upload .docx Template", type=["docx"], key="template")
    if template_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(template_file.getvalue())
            template_path = tmp.name
        st.sidebar.success("✅ Template loaded and ready.")
else:
    st.sidebar.markdown("**🔧 Select Formatting Options:**")
    formatting_choices = {
        "Bold Headings": st.sidebar.checkbox("Bold Headings", value=True),
        "Apply Page Breaks Before H1": st.sidebar.checkbox("Page Breaks Before Main Sections", value=True),
        "Standard Font": st.sidebar.selectbox("Font", ["Calibri", "Arial", "Times New Roman"]),
        "Font Size": st.sidebar.slider("Font Size", 10, 16, 11),
        "Justify Paragraphs": st.sidebar.checkbox("Justify Paragraphs", value=True),
        "Line Spacing": st.sidebar.selectbox("Line Spacing", ["Single", "1.15", "1.5", "Double"]),
        "Insert TOC": st.sidebar.checkbox("Insert TOC", value=True)
    }

def insert_toc(paragraph):
    def create_run_element(child_element):
        run = OxmlElement("w:r")
        run.append(child_element)
        return run

    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldSeparate = OxmlElement("w:fldChar")
    fldSeparate.set(qn("w:fldCharType"), "separate")

    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")

    paragraph._p.append(create_run_element(fldBegin))
    paragraph._p.append(create_run_element(instrText))
    paragraph._p.append(create_run_element(fldSeparate))
    paragraph._p.append(create_run_element(fldEnd))

def apply_formatting(doc_path, options):
    doc = Document(doc_path)

    spacing_map = {
        "Single": 1.0,
        "1.15": 1.15,
        "1.5": 1.5,
        "Double": 2.0
    }

    if options.get("Insert TOC"):
        toc_para = doc.paragraphs[0].insert_paragraph_before("📑 Table of Contents")
        insert_toc(toc_para)
        toc_para.style = "Normal"
        toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for para in doc.paragraphs:
        style_name = para.style.name.lower()

        if "heading 1" in style_name:
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_after = Pt(18)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if options.get("Bold Headings"):
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(14)
                para.text = para.text.upper()
            if options.get("Apply Page Breaks Before H1"):
                para.paragraph_format.page_break_before = True

        elif "heading 2" in style_name:
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if options.get("Bold Headings"):
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(12)
                    run.font.small_caps = True

        elif "heading" not in style_name:
            for run in para.runs:
                run.font.name = options.get("Standard Font", "Calibri")
                run.font.size = Pt(options.get("Font Size", 11))
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if options.get("Justify Paragraphs") else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = spacing_map.get(options.get("Line Spacing", "Single"))
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_together = True

    output_path = doc_path.replace(".docx", "_Formatted.docx")
    doc.save(output_path)

    if options.get("Insert TOC"):
        try:
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            docx = word.Documents.Open(output_path)
            docx.TablesOfContents(1).Update()
            docx.Save()
            docx.Close()
            word.Quit()
        except Exception as e:
            st.warning(f"⚠️ TOC update failed: {e}")

    return output_path

st.markdown('<div class="big-title">📘 TekSA Report Management System</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Upload a Word document to prepare it for professional formatting.</div>', unsafe_allow_html=True)

uploaded_file = st.file_uploader("📄 Upload a .docx file", type=["docx"])
temp_path = None

if uploaded_file:
    st.success("✅ File uploaded successfully.", icon="📂")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    st.download_button(
        label="📥 Download Unchanged Document",
        data=uploaded_file.getvalue(),
        file_name="TekSA_Original_Upload.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown(
        '<div class="info-box">ℹ️ This document has not been changed yet. Formatting will be applied in the next steps based on your selections in the sidebar.</div>',
        unsafe_allow_html=True
    )

if format_mode == "⚙️ Choose on-the-fly options" and formatting_choices:
    st.markdown("### 📋 Formatting Summary")
    st.markdown(
        f"""
        <div style='background-color:#f0f8ff;padding:15px 20px;border-left:6px solid #0077b6;border-radius:5px;'>
        <ul style="line-height:1.8">
        <li><strong>Font:</strong> {formatting_choices['Standard Font']}</li>
        <li><strong>Font Size:</strong> {formatting_choices['Font Size']} pt</li>
        {"<li><strong>Bold Headings:</strong> Enabled</li>" if formatting_choices['Bold Headings'] else ""}
        {"<li><strong>Page Breaks Before Headings:</strong> Enabled</li>" if formatting_choices['Apply Page Breaks Before H1'] else ""}
        {"<li><strong>Text Alignment:</strong> Justified</li>" if formatting_choices['Justify Paragraphs'] else "<li><strong>Text Alignment:</strong> Left</li>"}
        <li><strong>Line Spacing:</strong> {formatting_choices['Line Spacing']}</li>
        {"<li><strong>Insert Table of Contents:</strong> Yes</li>" if formatting_choices['Insert TOC'] else ""}
        </ul></div>
        """, unsafe_allow_html=True
    )

st.markdown("---")
st.subheader("📎 Final Step: Apply Formatting")

if st.button("✨ Apply Formatting"):
    if not uploaded_file:
        st.error("❌ Please upload a document to format.")
    elif format_mode == "📄 Upload a template" and not template_path:
        st.error("❌ Please upload a formatting template.")
    elif format_mode == "⚙️ Choose on-the-fly options":
        with st.spinner("🔧 Applying your selected formatting options..."):
            formatted_path = apply_formatting(temp_path, formatting_choices)
            with open(formatted_path, "rb") as f:
                st.download_button(
                    label="📥 Download Formatted Document",
                    data=f.read(),
                    file_name="TekSA_Formatted_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            st.success("✅ Formatting applied successfully!")
    else:
        st.warning("📄 Template formatting is not implemented yet.")
